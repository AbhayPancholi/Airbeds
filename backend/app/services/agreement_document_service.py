"""Generate Leave and Licence agreement document (filled template) as PDF or DOCX."""
from datetime import date, datetime
from pathlib import Path
import re
from typing import Any, Dict, List, Optional

from fastapi import HTTPException
from motor.motor_asyncio import AsyncIOMotorDatabase

from app.models.owner import OwnerResponse
from app.models.tenant import TenantResponse
from app.repositories.owner_repository import OwnerRepository
from app.repositories.tenant_repository import TenantRepository

# Page marker pattern: --page 1--, --page2--, etc. Content below each marker goes on one page.
PAGE_MARKER_RE = re.compile(r"--page\s*\d+\s*--", re.IGNORECASE)
MAIN_HEADING_TEXT = "LEAVE AND LICENCE AGREEMENT"

# Clause titles for bold prefix (spec: Clause Title in Bold, then body justified, first line indent 0.5")
CLAUSE_TITLES = (
    "Period", "License Fee & Deposit", "Payment of Deposit", "Maintenance Charges",
    "Electricity Charges", "Use", "Alteration", "No Tenancy", "Inspection", "Cancellation",
    "Possession", "Miscellaneous",
)
# Section headings that are centered and bold
CENTERED_HEADINGS = (
    "BETWEEN", "--AND---", "NOW THEREFORE THESE PRESENTS WITNESSETH THIS AGREEMENT AND IT IS HEREBY AGREED BY AND BETWEEN THE PARTIES HERETO AS FOLLOWS:",
    "SCHEDULE –A", "SCHEDULE -A",
    "IN WITNESS WHEREOF THE PARTIES HAVE SET AND SUBSCRIBED THEIR RESPECTIVE HANDS ON THE DAY AND THE YEAR FIRST HEREIN ABOVE MENTIONED",
)
# Left-aligned block labels (bold)
LEFT_HEADINGS = ("LICENSOR", "LICENSEE", "Signature:", "WITNESSES:")

# Style keys for PDF/DOCX
STYLE_TITLE = "title"                    # 14pt Bold Uppercase Center
STYLE_SECTION_CENTER = "section_center"  # 12pt Bold Center
STYLE_SECTION_LEFT = "section_left"      # 12pt Bold Left (LICENSOR, LICENSEE, etc.)
STYLE_BODY_LEFT = "body_left"            # 12pt Left (intro, party blocks)
STYLE_BODY_JUSTIFIED = "body_justified"  # 12pt Justified, first line indent 0.5" for clause body
STYLE_CLAUSE_TITLE = "clause_title"      # 12pt Bold (e.g. "Period:")
STYLE_CLAUSE_BODY = "clause_body"        # 12pt Justified, first line 0.5"
STYLE_MISC_ITEM = "misc_item"            # 12pt Justified (sub-points under Miscellaneous)
STYLE_BLANK = "blank"


def _split_into_pages(filled_text: str) -> List[str]:
    """Split text by --page N-- markers; each segment is one page. Strip markers from output."""
    parts = PAGE_MARKER_RE.split(filled_text)
    return [p.strip() for p in parts if p.strip()]


def _classify_line(line: str) -> tuple:
    """Return (style_key, display_text). For clauses may return two logical parts (title + body)."""
    s = line.strip()
    if not s:
        return [(STYLE_BLANK, "")]

    if s.upper() == MAIN_HEADING_TEXT.upper():
        return [(STYLE_TITLE, s.upper())]

    if s in CENTERED_HEADINGS:
        return [(STYLE_SECTION_CENTER, s)]
    if s.startswith("NOW THEREFORE") and "WITNESSETH" in s:
        return [(STYLE_SECTION_CENTER, s)]
    if s.startswith("IN WITNESS WHEREOF"):
        return [(STYLE_SECTION_CENTER, s)]
    if s == "--AND---":
        return [(STYLE_SECTION_CENTER, s)]

    if s in LEFT_HEADINGS:
        return [(STYLE_SECTION_LEFT, s)]

    if s.startswith("SCHEDULE") and "A" in s:
        return [(STYLE_SECTION_CENTER, s)]

    # Numbered clause: "1)\tPeriod: That the Licensor..."
    m = re.match(r"^(\d+)\)\s*([^:]+):\s*(.*)$", s)
    if m and m.group(2).strip() in CLAUSE_TITLES:
        title_part = m.group(2).strip() + ":"
        body_part = m.group(3).strip()
        if body_part:
            return [(STYLE_CLAUSE_TITLE, title_part), (STYLE_CLAUSE_BODY, body_part)]
        return [(STYLE_CLAUSE_TITLE, title_part)]

    # Clause without number: "12)\tThe Licensee shall..." (no colon title) - treat as body
    if re.match(r"^\d+\)\s+", s):
        return [(STYLE_CLAUSE_BODY, s)]

    # Miscellaneous sub-points: "1.\tAny existing damage..."
    if re.match(r"^\d+\.\s+", s) and "Miscellaneous" not in s:
        return [(STYLE_MISC_ITEM, s)]

    # Default: body justified for "Hereinafter", WHEREAS, etc.; body left for short lines like "1) Name:"
    if s.startswith("Hereinafter") or s.startswith("WHEREAS") or ("the Licensor" in s and "said premises" in s) or ("The Licensee has approached" in s):
        return [(STYLE_BODY_JUSTIFIED, s)]
    if s.startswith("1)") and "Name:" in s and ("LICENSOR" in s or "LICENSEE" in s or "P.O.A" in s or "tenant" in s.lower()):
        return [(STYLE_BODY_LEFT, s)]
    if s.startswith("THIS AGREEMENT") or s.startswith("On "):
        return [(STYLE_BODY_LEFT, s)]
    if "Flat No." in s and "admeasuring" in s and "SCHEDULE" not in s:
        return [(STYLE_BODY_JUSTIFIED, s)]
    if s.startswith("Flat No.") and "admeasuring" in s:
        return [(STYLE_BODY_JUSTIFIED, s)]
    if "Witnesses" in s or ("Name:" in s and "Add" in s):
        return [(STYLE_BODY_LEFT, s)]

    return [(STYLE_BODY_JUSTIFIED, s)]


def _iter_styled_blocks(filled_text: str):
    """Yield (style_key, text) for each paragraph; yield ("page_break", "") between pages."""
    pages = _split_into_pages(filled_text)
    for page_idx, page_text in enumerate(pages):
        for line in page_text.split("\n"):
            for style_key, text in _classify_line(line):
                if style_key == STYLE_BLANK:
                    yield (STYLE_BLANK, "")
                else:
                    yield (style_key, text)
        if page_idx < len(pages) - 1:
            yield ("page_break", "")


def _ordinal_suffix(n: int) -> str:
    if 10 <= n % 100 <= 20:
        return "th"
    return {1: "st", 2: "nd", 3: "rd"}.get(n % 10, "th")


def _age_from_dob(dob: Optional[str]) -> str:
    if not dob:
        return ""
    try:
        d = datetime.strptime(dob[:10], "%Y-%m-%d").date()
        return str((date.today() - d).days // 365)
    except (ValueError, TypeError):
        return ""


def _format_address(addr: Optional[Dict[str, Any]]) -> str:
    if not addr:
        return ""
    parts = [
        addr.get("flat_no"),
        addr.get("building_no"),
        addr.get("society"),
        addr.get("block_sector"),
        addr.get("street_landmark"),
        addr.get("city"),
        addr.get("state"),
        addr.get("pin_code"),
    ]
    return ", ".join(p for p in parts if p)


def _format_remaining_address(addr: Optional[Dict[str, Any]]) -> str:
    if not addr:
        return ""
    parts = [
        addr.get("society"),
        addr.get("block_sector"),
        addr.get("street_landmark"),
        addr.get("city"),
        addr.get("state"),
        addr.get("pin_code"),
    ]
    return ", ".join(p for p in parts if p)


class AgreementDocumentService:
    def __init__(self, db: AsyncIOMotorDatabase) -> None:
        self._owner_repo = OwnerRepository(db)
        self._tenant_repo = TenantRepository(db)
        self._template_path = Path(__file__).resolve().parent.parent / "templates" / "leave_and_licence_agreement.txt"

    def _load_template(self) -> str:
        if not self._template_path.exists():
            raise HTTPException(status_code=500, detail="Agreement template not found")
        return self._template_path.read_text(encoding="utf-8")

    async def build_filled_text(
        self,
        *,
        owner_id: str,
        tenant_id: str,
        flat_index: int = 0,
        agreement_date: str,
        start_date: str,
        end_date: str,
        monthly_rent: float,
        deposit_amount: float,
    ) -> str:
        owner = await self._owner_repo.get_by_id_response(owner_id)
        if not owner:
            raise HTTPException(status_code=404, detail="Owner not found")
        tenant = await self._tenant_repo.get_by_id(tenant_id)
        if not tenant:
            raise HTTPException(status_code=404, detail="Tenant not found")

        flats = owner.flats or []
        if flat_index < 0 or flat_index >= len(flats):
            raise HTTPException(status_code=400, detail="Invalid flat index")
        flat = flats[flat_index]
        flat_addr_raw = flat.address if hasattr(flat, "address") else (flat.get("address") if isinstance(flat, dict) else None)
        flat_addr = flat_addr_raw.model_dump() if flat_addr_raw and hasattr(flat_addr_raw, "model_dump") else (flat_addr_raw if isinstance(flat_addr_raw, dict) else {})

        # Agreement date parsing
        try:
            ag_d = datetime.strptime(agreement_date[:10], "%Y-%m-%d").date()
        except (ValueError, TypeError):
            ag_d = date.today()
        day_num = ag_d.day
        month_name = ag_d.strftime("%B")
        year_num = ag_d.year
        date_suffix = _ordinal_suffix(day_num)

        # Owner
        owner_addr = (owner.address.model_dump() if hasattr(owner.address, "model_dump") else owner.address) if owner.address else {}
        owner_sal = ""  # owner has no salutation in model
        owner_name = owner.name or ""
        owner_age = _age_from_dob(owner.dob)
        owner_occupation = owner.occupation or ""
        owner_flat_no = owner_addr.get("flat_no", "")
        owner_building = owner_addr.get("building_no", "")
        owner_block = owner_addr.get("block_sector", "")
        owner_remaining = ", ".join(
            p
            for p in [
                owner_addr.get("street_landmark"),
                owner_addr.get("city"),
                owner_addr.get("state"),
                owner_addr.get("pin_code"),
            ]
            if p
        )

        # POA
        poa = owner.poa
        if poa:
            poa_addr = (poa.address.model_dump() if hasattr(poa.address, "model_dump") else poa.address) if poa.address else {}
            poa_sal = ""
            poa_name = poa.name or ""
            poa_age = _age_from_dob(poa.dob)
            poa_occupation = poa.occupation or ""
            poa_full_addr = _format_address(poa_addr)
        else:
            poa_sal = poa_name = poa_age = poa_occupation = poa_full_addr = ""

        # Tenant
        tenant_sal = (tenant.salutation or "").strip()
        tenant_name = tenant.tenant_name or ""
        tenant_age = tenant.age or _age_from_dob(tenant.dob)
        tenant_occupation = (tenant.employment_type or tenant.occupancy_details or "").strip()
        tenant_full_addr = ", ".join(
            p for p in [tenant.residential_address, tenant.pin_code, tenant.state, tenant.country] if p
        )

        # Flat (rented premises)
        flat_building = flat_addr.get("building_no", "") or flat_addr.get("building_name", "")
        flat_no = flat_addr.get("flat_no", "")
        area_val = getattr(flat, "measurement_sqft", None) or (flat.get("measurement_sqft") if isinstance(flat, dict) else None)
        area_str = str(int(area_val)) if area_val is not None else ""
        floor_no = getattr(flat, "floor_no", None) or (flat.get("floor_no") if isinstance(flat, dict) else None) or ""
        floor_no_str = str(floor_no)
        floor_suffix = _ordinal_suffix(int(floor_no)) if floor_no_str.isdigit() else ""
        bhk = getattr(flat, "bhk", None) or (flat.get("bhk") if isinstance(flat, dict) else None) or ""
        car_parking = getattr(flat, "car_parking", None)
        if car_parking is None and isinstance(flat, dict):
            car_parking = flat.get("car_parking", False)
        car_parking = bool(car_parking)
        with_without = "with car parking" if car_parking else "without car parking"
        schedule_car = "with One Covered Car Parking on Ground Floor" if car_parking else "without Car Parking"
        flat_remaining = _format_remaining_address(flat_addr)
        flat_full_addr = _format_address(flat_addr)

        replacements = {
            "[Date]": str(day_num),
            "[Date suffix]": date_suffix,
            "[Month]": month_name,
            "[Year]": str(year_num),
            "[Owner salutation]": owner_sal,
            "[Name of the flat owner]": owner_name,
            "[full name of the owner]": owner_name,
            "[Age of the owner]": owner_age,
            "[occupation of owner]": owner_occupation,
            "[Flat No of owner]": owner_flat_no,
            "[Building Name]": owner_building,
            "[Block/Sector]": owner_block,
            "[Remaining address with PIN code]": owner_remaining,
            "[POA salutation]": poa_sal,
            "[Full name of the POA]": poa_name,
            "[Name of POA]": poa_name,
            "[Age of POA]": poa_age,
            "[occupation of POA]": poa_occupation,
            "[Full Address of POA]": poa_full_addr,
            "[Salutation]": tenant_sal,
            "[Full name of the Tenant]": tenant_name,
            "[full name of the tenant]": tenant_name,
            "[Age of Tenant]": tenant_age,
            "[Occupation of Tenant]": tenant_occupation,
            "[Full Address of Tenant]": tenant_full_addr,
            "[building name of the flat rent by owner]": flat_building,
            "[flat number]": flat_no,
            "[Area]": area_str,
            "[area of the owner's rented flat]": area_str,
            "[Area of the rented flat]": area_str,
            "[Floor number]": floor_no_str,
            "[Floor suffix]": floor_suffix,
            "[floor number]": floor_no_str,
            "[size]": bhk,
            "[with or without car parking]": with_without,
            "[Schedule car parking]": schedule_car,
            "[remaining address of the owner's rented flat like society, block, street, city, state, and pincode]": flat_remaining,
            "[remaining full address of the rented flat]": flat_full_addr,
            "[owner's rented flat with building name]": f"{flat_building} {flat_no}".strip(),
            "[start-date]": start_date[:10] if start_date else "",
            "[end-date]": end_date[:10] if end_date else "",
            "[monthly rent for the owner's rented flat]": str(int(monthly_rent)),
            "[One time deposit]": str(int(deposit_amount)),
        }

        template = self._load_template()
        for key, value in replacements.items():
            template = template.replace(key, value or "___")
        return template

    def to_pdf_bytes(self, filled_text: str) -> bytes:
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import inch
        from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

        buffer = __import__("io").BytesIO()
        doc = SimpleDocTemplate(
            buffer, pagesize=A4,
            rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=72,
        )
        base = getSampleStyleSheet()["Normal"]
        # Times-Roman is built-in in PDF; Bold variant is Times-Bold
        fn, fn_bold = "Times-Roman", "Times-Bold"

        def _para_style(**kwargs):
            opts = {"name": "", "parent": base, "fontName": fn, "fontSize": 12, "leading": 12 * 1.15}
            opts.update(kwargs)
            return ParagraphStyle(**opts)

        styles = {
            STYLE_TITLE: ParagraphStyle(
                name="", parent=base, fontName=fn_bold, fontSize=14, alignment=1,
                spaceAfter=6, leading=14 * 1.15,
            ),
            STYLE_SECTION_CENTER: _para_style(fontName=fn_bold, alignment=1, spaceAfter=6),
            STYLE_SECTION_LEFT: _para_style(fontName=fn_bold, alignment=0, spaceAfter=6),
            STYLE_BODY_LEFT: _para_style(alignment=0, spaceAfter=6),
            STYLE_BODY_JUSTIFIED: _para_style(alignment=4, spaceAfter=6),  # 4 = justify
            STYLE_CLAUSE_TITLE: _para_style(fontName=fn_bold, alignment=0, spaceAfter=2),
            STYLE_CLAUSE_BODY: _para_style(alignment=4, firstLineIndent=0.5 * inch, spaceAfter=6),
            STYLE_MISC_ITEM: _para_style(alignment=4, leftIndent=0.25 * inch, spaceAfter=4),
            STYLE_BLANK: None,
        }
        story = []
        for style_key, text in _iter_styled_blocks(filled_text):
            if style_key == "page_break":
                story.append(PageBreak())
                continue
            if style_key == STYLE_BLANK:
                story.append(Spacer(1, 6))
                continue
            esc = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
            style = styles.get(style_key) or styles[STYLE_BODY_JUSTIFIED]
            story.append(Paragraph(esc, style))
        doc.build(story)
        return buffer.getvalue()

    def to_docx_bytes(self, filled_text: str) -> bytes:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        buffer = __import__("io").BytesIO()
        doc = Document()
        # Set default font Times New Roman, 12pt, line spacing 1.15
        normal_style = doc.styles["Normal"]
        normal_style.font.name = "Times New Roman"
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.space_after = Pt(6)
        normal_style.paragraph_format.line_spacing = 1.15

        for style_key, text in _iter_styled_blocks(filled_text):
            if style_key == "page_break":
                doc.add_page_break()
                continue
            if style_key == STYLE_BLANK:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(6)
                continue
            p = doc.add_paragraph()
            run = p.add_run(text)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15
            if style_key == STYLE_TITLE:
                run.bold = True
                run.font.size = Pt(14)
                run.font.all_caps = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif style_key == STYLE_SECTION_CENTER:
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif style_key == STYLE_SECTION_LEFT:
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif style_key == STYLE_BODY_LEFT:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif style_key in (STYLE_BODY_JUSTIFIED, STYLE_CLAUSE_BODY, STYLE_MISC_ITEM):
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            elif style_key == STYLE_CLAUSE_TITLE:
                run.bold = True
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_after = Pt(2)
            if style_key == STYLE_CLAUSE_BODY:
                p.paragraph_format.first_line_indent = Inches(0.5)
            if style_key == STYLE_MISC_ITEM:
                p.paragraph_format.left_indent = Inches(0.25)
                p.paragraph_format.space_after = Pt(4)
        doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
